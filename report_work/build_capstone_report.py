from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path.cwd()
WORK = ROOT / "report_work"
TEMPLATE = WORK / "template.docx"
OUT = WORK / "Competra_Final_Capstone_Report_Fred_Niyonshuti.docx"
SCREENSHOTS = ROOT / "docs" / "screenshots"
DIAGRAMS = WORK / "generated_diagrams"
DIAGRAMS.mkdir(exist_ok=True)

TITLE = "Design and Development of a Skills Gap Analysis Tool for ICT TVET Graduates in Kicukiro District, Rwanda"
STUDENT = "Fred Niyonshuti"
PROGRAM = "Bachelor of Software Engineering"
INSTITUTION = "African Leadership University"
COURSE = "Final-Year Software Engineering Capstone Project"
YEAR = "2026"
SUPERVISOR = "Supervisor Name"

PRIMARY = RGBColor(0x00, 0x77, 0xB6)
DARK = RGBColor(0x03, 0x04, 0x5E)
TEAL = RGBColor(0x0F, 0x52, 0x57)
MUTED = RGBColor(0x4B, 0x5D, 0x75)


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if not child.tag.endswith("sectPr"):
            body.remove(child)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    doc.styles["Normal"].font.size = Pt(10.5)


def paragraph(doc, text="", bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.bold = bold
    r.italic = italic
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.bold = True
    r.font.color.rgb = DARK if level == 1 else TEAL
    r.font.size = Pt(15 if level == 1 else 12.5 if level == 2 else 11)
    return p


_number_counter = 0


def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    r = p.add_run("- " + text)
    r.font.name = "Arial"
    r.font.size = Pt(10.2)
    return p


def numbered(doc, text):
    global _number_counter
    _number_counter += 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    r = p.add_run(f"{_number_counter}. {text}")
    r.font.name = "Arial"
    r.font.size = Pt(10.2)
    return p


def caption(doc, text):
    p = paragraph(doc, text, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = MUTED


def set_cell(cell, text, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.bold = header
    if header:
        r.font.color.rgb = RGBColor(255, 255, 255)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0077B6")
        tc_pr.append(shd)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, title=None):
    if title:
        caption(doc, title)
    t = doc.add_table(rows=1, cols=len(headers))
    if 'Table Grid' in [style.name for style in doc.styles]:
        t.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph("")
    return t


def add_image(doc, image_path, title, width=6.3):
    image_path = Path(image_path)
    if not image_path.exists():
        paragraph(doc, f"[Missing image: {image_path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    caption(doc, title)


def add_toc(doc):
    heading(doc, "Table of Contents", 1)
    paragraph(
        doc,
        "Right-click this field in Microsoft Word and select Update Field if page numbers are not displayed automatically.",
        italic=True,
    )
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)
    doc.add_page_break()


def font(size=24, bold=False):
    file = "arialbd.ttf" if bold else "arial.ttf"
    path = Path("C:/Windows/Fonts") / file
    if path.exists():
        return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap(draw, text, fnt, width):
    words, lines, current = text.split(), [], ""
    for word in words:
        attempt = (current + " " + word).strip()
        if draw.textbbox((0, 0), attempt, font=fnt)[2] <= width:
            current = attempt
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def box(draw, xy, text, fill="#FFFFFF"):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline="#0077B6", width=3)
    x1, y1, x2, y2 = xy
    fnt = font(22, True)
    lines = wrap(draw, text, fnt, x2 - x1 - 30)
    y = y1 + ((y2 - y1) - len(lines) * 28) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, fill="#03045E", font=fnt)
        y += 30


def arrow(draw, start, end):
    draw.line([start, end], fill="#0F5257", width=4)
    x, y = end
    draw.polygon([(x, y), (x - 14, y - 8), (x - 14, y + 8)], fill="#0F5257")


def create_diagram(name, title, items):
    img = Image.new("RGB", (1500, 850), "#F7FAFC")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), title, fill="#03045E", font=font(34, True))
    y = 140
    for index, item in enumerate(items, start=1):
        box(draw, (260, y, 1240, y + 85), f"{index}. {item}")
        if index < len(items):
            draw.line((750, y + 85, 750, y + 125), fill="#0F5257", width=4)
        y += 125
    path = DIAGRAMS / f"{name}.png"
    img.save(path)
    return path


def create_architecture():
    img = Image.new("RGB", (1500, 900), "#F7FAFC")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "Competra Full Stack Architecture", fill="#03045E", font=font(34, True))
    boxes = [
        ("React + Vite Frontend", (80, 150, 430, 260)),
        ("Express API", (575, 150, 925, 260)),
        ("Services Layer", (575, 380, 925, 490)),
        ("MongoDB Atlas", (1070, 380, 1420, 490)),
        ("GitHub, E2B, Gemini, Brevo, Google OAuth", (180, 650, 1320, 760)),
    ]
    for text, xy in boxes:
        box(draw, xy, text, "#E8F1F8")
    for start, end in [((430, 205), (575, 205)), ((750, 260), (750, 380)), ((925, 435), (1070, 435)), ((750, 490), (750, 650))]:
        arrow(draw, start, end)
    path = DIAGRAMS / "architecture.png"
    img.save(path)
    return path


def preliminaries(doc):
    paragraph(doc, "")
    paragraph(doc, "")
    paragraph(doc, TITLE, True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(19)
    paragraph(doc, "(Competra)", True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.color.rgb = PRIMARY
    for line in [f"Prepared by: {STUDENT}", f"Program: {PROGRAM}", f"Course: {COURSE}", f"Supervisor: {SUPERVISOR}", f"Institution: {INSTITUTION}", f"Year: {YEAR}"]:
        paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    heading(doc, "DECLARATION", 1)
    paragraph(doc, "This report was my original work unless stated otherwise, and all external sources were cited and referenced. The work had not been presented for the award of a degree or for a similar purpose in any other university.")
    paragraph(doc, "Signature: ________________________________        Date: ____________________")
    paragraph(doc, f"Name of Student: {STUDENT}")
    doc.add_page_break()

    heading(doc, "CERTIFICATION", 1)
    paragraph(doc, f"The undersigned certified that the report entitled \"{TITLE}\" had been reviewed and was recommended for acceptance by {INSTITUTION}.")
    paragraph(doc, "Signature: ________________________________        Date: ____________________")
    paragraph(doc, f"{SUPERVISOR}")
    paragraph(doc, f"Faculty, {PROGRAM}, ALU")
    doc.add_page_break()

    heading(doc, "DEDICATION AND ACKNOWLEDGEMENT", 1)
    paragraph(doc, "This report was dedicated to my family, lecturers, supervisor, classmates, and ICT TVET stakeholders who supported the development and review of this capstone project.")
    paragraph(doc, "I acknowledged the guidance received from African Leadership University and the feedback that shaped the design, implementation, testing, and deployment of the Competra system.")
    doc.add_page_break()

    heading(doc, "Abstract", 1)
    paragraph(doc, "This project designed and developed Competra, a web-based Skills Gap Analysis Tool for ICT TVET graduates in Kicukiro District, Rwanda. The study responded to the lack of structured tools that compared graduate ICT competencies with RTB-aligned requirements. The system was implemented using React, Vite, TypeScript, Node.js, Express.js, MongoDB, JWT authentication, Google OAuth, GitHub repository review, E2B sandbox execution, Gemini recommendations, and Brevo transactional email. The workflow allowed learners to select a competency, submit GitHub practical evidence and theory answers, receive automatic practical and theory scores, compare the final score with a benchmark, view a classified gap result, receive recommendations with resources, and download reports. Testing used unit, integration, validation, functional, responsive, deployment, and external API strategies. Backend tests passed 48 out of 48 checks, frontend tests passed 10 out of 10 checks, and the production frontend build completed successfully. The results showed that the system implemented the core features required by the proposal, although large-scale field evaluation data from 50-120 graduates had not yet been provided.")
    doc.add_page_break()
    add_toc(doc)

    heading(doc, "List of Tables", 1)
    for item in ["Table 1.1: Research Budget", "Table 1.2: Research Timeline", "Table 2.1: Summary of Reviewed Literature", "Table 3.1: Functional Requirements", "Table 3.2: Non-Functional Requirements", "Table 3.3: Database Model Summary", "Table 4.1: Testing Results Summary", "Table 5.1: Objective Achievement Analysis"]:
        paragraph(doc, item)
    doc.add_page_break()
    heading(doc, "List of Figures", 1)
    for item in ["Figure 3.1: Full Stack System Architecture Diagram", "Figure 3.2: Core Assessment Workflow Diagram", "Figure 3.3: Use Case Diagram", "Figure 4.1: Home Page", "Figure 4.2: Sign In Page", "Figure 4.3: Create Account Page", "Figure 4.4: Learner Dashboard", "Figure 4.5: Admin Dashboard", "Figure 4.6: Competency Management", "Figure 4.7: Take Assessment Page", "Figure 4.8: Gap Results", "Figure 4.9: Recommendations", "Figure 4.10: Reports", "Figure 4.11: Notifications", "Figure 4.12: Testing and Deployment Evidence"]:
        paragraph(doc, item)
    doc.add_page_break()
    heading(doc, "List of Acronyms/Abbreviations", 1)
    table(doc, ["Acronym", "Meaning"], [("API", "Application Programming Interface"), ("CBET", "Competency-Based Education and Training"), ("ICT", "Information and Communication Technology"), ("JWT", "JSON Web Token"), ("RTB", "Rwanda TVET Board"), ("TVET", "Technical and Vocational Education and Training"), ("UI", "User Interface"), ("UX", "User Experience"), ("UML", "Unified Modelling Language")])
    doc.add_page_break()


def chapter_one(doc):
    heading(doc, "CHAPTER ONE: INTRODUCTION", 1)
    heading(doc, "1.1 Introduction and Background", 2)
    paragraph(doc, "Technical and Vocational Education and Training was treated as an important pathway for preparing skilled workers in Rwanda. The approved proposal showed that Rwanda's TVET system emphasized competency-based training and occupational standards through RTB-aligned frameworks. Rwanda's ICT sector also required practical digital skills because national digital-skills strategies and labour-market reports had linked ICT skills to employability and transformation (Rwanda Polytechnic, 2023; Rwanda Information Society Authority, 2023; Rwanda Development Board, 2022).")
    paragraph(doc, "This project therefore focused on ICT TVET graduates and built a practical system that compared demonstrated competence with benchmark expectations. The system did not rely only on self-reported ability. It used GitHub repository evidence, theory answers, benchmark comparison, gap classification, AI-supported recommendations, notifications, and reports.")
    heading(doc, "1.2 Problem statement", 2)
    paragraph(doc, "The problem addressed by this project was the absence of a structured web-based tool that assessed ICT TVET graduates against RTB-aligned competency requirements and explained their specific skill gaps. Existing approaches were often manual, informal, or institution-specific, which made it difficult for learners to identify strengths, weaknesses, and clear improvement actions.")
    heading(doc, "1.3 Project's main objective", 2)
    paragraph(doc, "The main objective was to design and develop a web-based Skills Gap Analysis Tool that analysed ICT TVET graduates' competencies against Rwanda TVET Board occupational standards in order to identify skills gaps and improve employability alignment within Kicukiro District, Rwanda.")
    heading(doc, "1.3.1 List of the specific objectives", 3)
    for text in ["To assess ICT TVET graduate competency levels using practical and theory evidence.", "To design and develop a web-based tool that mapped competency scores against benchmark requirements.", "To test and evaluate the effectiveness, usability, and technical accuracy of the developed system.", "To support competency awareness through gap results, recommendations, notifications, and reports."]:
        numbered(doc, text)
    heading(doc, "1.4 Research questions", 2)
    for text in ["What challenges did ICT TVET graduates in Kicukiro District face in aligning their competencies with RTB occupational standards?", "What limitations existed in current ICT competency assessment methods?", "What functional and non-functional requirements were needed for an effective Skills Gap Analysis Tool?", "How effective was the developed system in identifying competency gaps and supporting employability alignment?"]:
        numbered(doc, text)
    heading(doc, "1.5 Project scope", 2)
    paragraph(doc, "The project was limited to ICT-related TVET competencies in Kicukiro District. The implemented system supported learners, organization users, organization administrators, administrators, and super administrators. It included account creation, email verification, authentication, organization management, competency management, benchmark management, checklist management, GitHub repository review, theory scoring, gap analysis, recommendations, notifications, reports, dashboards, and deployment.")
    heading(doc, "1.6 Significance and Justification", 2)
    paragraph(doc, "The system was significant because it helped convert practical evidence into measurable feedback. Learners received gap levels and resource-based recommendations, while administrators obtained tools for managing competencies, benchmarks, and assessment data. This supported literature that emphasized practical competence, employability, and institutional TVET quality improvement (Bhorat et al., 2023; UNESCO-UNEVOC, 2017).")
    heading(doc, "1.7 Research Budget", 2)
    table(doc, ["Item", "Purpose", "Estimated Cost (RWF)"], [("Internet and hosting", "Development, API testing, and deployment", "120,000"), ("Software tools", "Design, coding, and documentation", "50,000"), ("Data collection", "Questionnaires and stakeholder feedback", "80,000"), ("Contingency", "Unexpected project expenses", "50,000"), ("Total", "Estimated budget", "300,000")], "Table 1.1: Research Budget")
    heading(doc, "1.8 Research Timeline", 2)
    table(doc, ["Activity", "Period", "Status"], [("Requirement analysis and proposal review", "Month 1", "Completed"), ("System design and database modelling", "Month 1-2", "Completed"), ("Backend and frontend implementation", "Month 2-3", "Completed"), ("Testing, deployment, and report writing", "Month 3-4", "Completed; broad field evaluation evidence was pending")], "Table 1.2: Research Timeline")


def chapter_two(doc):
    heading(doc, "CHAPTER TWO: LITERATURE REVIEW", 1)
    heading(doc, "2.1 Introduction", 2)
    paragraph(doc, "This chapter reviewed literature related to TVET competency assessment, digital-skills gaps, web-based assessment systems, and the technologies used to implement evidence-based skills-gap analysis.")
    heading(doc, "2.2 Historical Background of the Research Topic", 2)
    paragraph(doc, "Competency-based education and training had increasingly focused on demonstrable skills. Rwanda's TVET framework emphasized occupational standards and practical assessment, while digital-skills frameworks showed that ICT competence had become central to employability and national transformation (Rwanda Polytechnic, 2023; Rwanda Information Society Authority, 2023).")
    heading(doc, "2.3 Overview of Existing System", 2)
    paragraph(doc, "Existing approaches included manual assessment records, spreadsheets, learning management systems, and instructor feedback. These approaches helped manage some records but often did not automatically score practical repository evidence, compare results with benchmarks, classify gaps, and provide personalized improvement resources.")
    heading(doc, "2.4 Review of Related Work", 2)
    table(doc, ["Source", "Focus", "Relevance"], [("Rwanda Polytechnic (2023)", "TVET CBC framework", "Supported occupational standards and competency assessment."), ("RISA (2023)", "National digital skills", "Supported the need for ICT skill tracking."), ("RDB (2022)", "ICT labour market", "Explained sector skills and employment needs."), ("Bhorat et al. (2023)", "Digital skills gaps in Africa", "Supported skills-gap problem context."), ("Hakizayezu and Maniraho (2022)", "TVET and youth employment", "Supported employability-alignment challenges."), ("UNESCO-UNEVOC (2017)", "TVET provider toolkits", "Supported quality-improvement practices.")], "Table 2.1: Summary of Reviewed Literature")
    heading(doc, "2.5 Strength and Weaknesses of the Existing System(s)", 2)
    paragraph(doc, "The reviewed approaches were useful for storing records and supporting manual assessment, but they had weaknesses in automatic evidence review, benchmark comparison, AI-supported recommendations, learner-friendly reports, and centralized analytics. Competra addressed these weaknesses through a full-stack system that integrated assessment, scoring, recommendation, and reporting features.")
    heading(doc, "2.6 General Comments", 2)
    paragraph(doc, "The literature confirmed that a realistic ICT skills-gap system needed both education alignment and technical correctness. Therefore, the implemented system combined competency-based assessment principles with a secure web architecture.")


def chapter_three(doc, architecture, workflow, usecase):
    heading(doc, "CHAPTER THREE: SYSTEM ANALYSIS AND DESIGN", 1)
    heading(doc, "3.1 Introduction", 2)
    paragraph(doc, "This chapter described the design that guided the completed implementation, including the research design, requirements, architecture, database design, UML diagrams, and development tools.")
    heading(doc, "3.2 Research Design (including the SDLC model used)", 2)
    paragraph(doc, "The project followed an Agile development model. This model was selected because requirements improved through repeated feedback and implementation cycles. The system was designed, implemented, tested, deployed, and refined through iterations that improved authentication, dashboards, repository review, scoring, recommendation generation, UI responsiveness, and deployment.")
    heading(doc, "3.3 Functional and Non-functional Requirements", 2)
    table(doc, ["Functional Requirement", "Implemented Status"], [("Authentication and authorization", "Implemented with JWT, Google OAuth, email verification, password reset, and role middleware."), ("Graduate profile management", "Implemented for learner details and organization selection."), ("Competency, checklist, and benchmark management", "Implemented for administrators."), ("Assessment submission", "Implemented using GitHub repository evidence and theory answers."), ("Skills-gap analysis", "Implemented with weighted scoring and benchmark comparison."), ("Recommendations, reports, and notifications", "Implemented using Gemini, MongoDB, in-app notifications, email support, and downloadable reports.")], "Table 3.1: Functional Requirements")
    table(doc, ["Non-Functional Requirement", "Implementation Approach"], [("Security", "Password hashing, JWT, role checks, rate limiting, email verification, and environment variables."), ("Usability", "Responsive dashboards, tables, modals, charts, filters, and clear messages."), ("Maintainability", "Controller-service-model backend structure and reusable frontend components."), ("Reliability", "Automated tests, validation, structured errors, and deployment health checks.")], "Table 3.2: Non-Functional Requirements")
    heading(doc, "3.4 System Architecture", 2)
    paragraph(doc, "The system used React and Vite on the frontend, Express.js on the backend, MongoDB Atlas for persistence, and external services for GitHub, E2B, Gemini, Google OAuth, and Brevo email. The backend followed a layered structure from server.js to app.js, routes, middleware, controllers, services, models, and the database.")
    add_image(doc, architecture, "Figure 3.1: Full Stack System Architecture Diagram")
    heading(doc, "3.5 Flow Chart, Class diagram, Use Case Diagram, Sequence Diagram and all other diagrams.", 2)
    add_image(doc, workflow, "Figure 3.2: Core Assessment Workflow Diagram")
    add_image(doc, usecase, "Figure 3.3: Use Case Diagram")
    heading(doc, "3.6 Database Design", 2)
    table(doc, ["Model", "Purpose"], [("User", "Stored accounts, roles, organization, password hash, email verification, and reset tokens."), ("Organization", "Stored TVET institutions and organization records."), ("Competency", "Stored ICT competency code, category, practical tasks, and theory questions."), ("Checklist", "Stored admin-defined repository checklist items and weights."), ("Benchmark", "Stored RTB-aligned required score per competency."), ("Assessment", "Stored evidence, scores, benchmark, skill gap, and review status."), ("RepositoryAssessmentResult", "Stored repository execution and review evidence."), ("Recommendation", "Stored Gemini recommendation, action items, and resources."), ("Report and Notification", "Stored generated reports and user notifications.")], "Table 3.3: Database Model Summary")
    heading(doc, "3.7 Development Tools", 2)
    table(doc, ["Tool/Technology", "Purpose"], [("React, Vite, TypeScript, Tailwind CSS", "Frontend user interface."), ("Node.js and Express.js", "Backend API and business logic."), ("MongoDB Atlas and Mongoose", "Database persistence."), ("GitHub API and E2B", "Repository verification and isolated execution."), ("Gemini API", "Recommendation generation."), ("Brevo API", "Email verification and password reset."), ("Vitest and Node test runner", "Automated testing."), ("Vercel and Render", "Frontend and backend deployment.")], "Table 3.4: Development Tools and Technologies")


def chapter_four(doc):
    heading(doc, "CHAPTER 4: SYSTEM IMPLEMENTATION AND TESTING", 1)
    heading(doc, "4.1 Implementation and coding", 2)
    heading(doc, "4.1.1 Introduction", 3)
    paragraph(doc, "This section described the implemented software modules, screenshots, source-code evidence, and testing outputs. It focused on the features that directly addressed the problem statement.")
    heading(doc, "4.1.2 Description of implementation Tools and technology", 3)
    paragraph(doc, "The frontend was implemented using React, Vite, TypeScript, Tailwind CSS, Recharts, and Lucide icons. The backend was implemented using Node.js, Express.js, ES Modules, MongoDB, Mongoose, JWT, Google OAuth, GitHub API, E2B, Gemini API, and Brevo email API.")
    heading(doc, "4.2 Graphical view of the project", 2)
    for name, title in [("homepage.png", "Figure 4.1: Home Page"), ("sign-in.png", "Figure 4.2: Sign In Page"), ("sign-up.png", "Figure 4.3: Create Account Page"), ("user-dashboard.png", "Figure 4.4: Learner Dashboard"), ("admin-dashboard.png", "Figure 4.5: Admin Dashboard"), ("competency.png", "Figure 4.6: Competency Management"), ("assessment.png", "Figure 4.7: Take Assessment Page"), ("gapresult.png", "Figure 4.8: Skill Gap Results"), ("recommendation.png", "Figure 4.9: Recommendations"), ("report.png", "Figure 4.10: Reports"), ("notification.png", "Figure 4.11: Notifications")]:
        add_image(doc, SCREENSHOTS / name, title)
        paragraph(doc, f"{title.replace('Figure', 'The figure')} showed a completed system feature that supported the planned functionality.")
    heading(doc, "4.3 Testing", 2)
    paragraph(doc, "Testing was performed using unit, integration, validation, functional, responsive, deployment, and external integration testing strategies.")
    table(doc, ["Testing Strategy", "Observed Result", "Status"], [("Backend unit and integration testing", "48 tests passed out of 48.", "Passed"), ("Frontend component and utility testing", "10 tests passed out of 10 across 3 files.", "Passed"), ("Backend build check", "node --check src/server.js completed.", "Passed"), ("Frontend production build", "Vite build completed successfully.", "Passed with bundle-size warning"), ("Validation testing", "Invalid emails, weak passwords, invalid scores, and invalid checklist totals were rejected.", "Passed"), ("Deployment testing", "Backend health endpoint and deployed frontend were verified.", "Passed")], "Table 4.1: Testing Results Summary")
    for name, title in [("backendtest.png", "Figure 4.12: Backend Testing Evidence"), ("frontendtest.png", "Figure 4.13: Frontend Testing Evidence"), ("validation2.png", "Figure 4.14: Validation Testing Evidence"), ("performance.png", "Figure 4.15: Performance Testing Evidence"), ("responsive.png", "Figure 4.16: Responsive Testing Evidence"), ("frontend-deployment.png", "Figure 4.17: Frontend Deployment Evidence"), ("backend-deployment.png", "Figure 4.18: Backend Deployment Evidence"), ("application-production.png", "Figure 4.19: Production Application Evidence")]:
        add_image(doc, SCREENSHOTS / name, title)
    paragraph(doc, "Acceptance testing evidence showed that the core workflow was usable from competency selection to report viewing. However, final acceptance survey data from the proposed 50-120 ICT TVET graduates had not yet been provided.")


def chapter_five(doc):
    heading(doc, "CHAPTER 5: THE DESCRIPTION OF THE RESULTS/SYSTEM", 1)
    paragraph(doc, "The final system produced evidence-based competency results. A learner selected a competency, submitted a GitHub repository and theory answers, and received an automatically reviewed result. The assessment service calculated the final score using 70% GitHub practical task score and 30% theory score. The system compared this final score with the benchmark score, calculated the skill gap, classified the gap level, generated a Gemini recommendation, notified the learner, and generated a report.")
    heading(doc, "5.1 Repository Review Result", 2)
    paragraph(doc, "The repository review was stored in MongoDB with repository metadata, sampled source evidence, checklist results, E2B execution evidence where configured, ESLint result, security result, passed requirements, failed requirements, competency scores, and improvement recommendations. Admin-created checklist weights controlled practical scoring, which made the result more aligned with the selected task.")
    heading(doc, "5.2 Recommendation Result", 2)
    paragraph(doc, "Gemini recommendations were generated from measured scores, benchmark score, skill gap, gap level, failed repository checks, strengths, weak areas, and repository summaries. The recommendation service required learning resource links before saving a recommendation, so the user received both guidance and resources for closing the gap.")
    heading(doc, "5.3 Analysis of Results", 2)
    table(doc, ["Objective", "Achievement", "Analysis"], [("Assess ICT TVET competency levels", "Partially achieved", "The system assessed repository evidence and theory answers; broad pilot data was still missing."), ("Develop a web-based skills-gap tool", "Achieved", "Core MERN-style application, dashboards, authentication, scoring, recommendations, reports, and deployment were implemented."), ("Test effectiveness, usability, and accuracy", "Partially achieved", "Automated tests and screenshots were available; full user survey and expert agreement results were pending."), ("Improve competency awareness", "Partially achieved", "The system displayed scores, gaps, recommendations, resources, notifications, and reports; before/after awareness survey results were pending.")], "Table 5.1: Objective Achievement Analysis")
    heading(doc, "5.4 Discussion", 2)
    paragraph(doc, "The milestones showed that a practical skills-gap system could be implemented using objective evidence rather than self-report only. The system's impact was that learners could identify the exact competency, score, benchmark, gap, weak practical checks, and improvement resources. The limitation was that repository automation could not perfectly judge every programming language or coding style without well-designed task contracts, checklist rules, hidden tests, and assessor validation.")


def chapter_six(doc):
    heading(doc, "CHAPTER 6: CONCLUSIONS AND RECOMMENDATIONS", 1)
    heading(doc, "6.1 Conclusion", 2)
    paragraph(doc, "The project successfully designed and implemented Competra as a web-based Skills Gap Analysis Tool for ICT TVET graduates in Kicukiro District. The implemented system addressed the problem by collecting practical GitHub evidence and theory answers, calculating weighted competency scores, comparing them with RTB-aligned benchmarks, classifying skill gaps, generating Gemini recommendations with resources, notifying users, and generating reports.")
    paragraph(doc, "The system was technically aligned with the approved proposal and was verified through automated backend tests, frontend tests, validation checks, production build checks, deployment evidence, and screenshots. The main unresolved limitation was the absence of full field-evaluation evidence from the proposed graduate sample.")
    heading(doc, "6.2 Recommendations", 2)
    for item in ["TVET institutions should use evidence-based digital assessment tools to complement manual assessment.", "ICT graduates should submit repositories with clear README files, tests, and deployment instructions.", "Administrators should define measurable checklist requirements and benchmark scores before repository review.", "Future development should integrate official RTB data, employer dashboards, plagiarism checks, and broader language support."]:
        bullet(doc, item)
    heading(doc, "6.3 Suggestions for Further Studies or Research", 2)
    paragraph(doc, "Further studies should compare Gemini recommendations with ICT instructor recommendations, measure before-and-after competency awareness, and evaluate whether benchmark-based reports improved employability readiness among ICT TVET graduates.")


def references(doc):
    heading(doc, "References", 1)
    refs = [
        "Bhorat, H., Signe, L., Asmal, Z., Monnakgotla, J., & Rooney, C. (2023). Digitalisation and digital skills gaps in Africa. Brookings Institution. https://www.brookings.edu/wp-content/uploads/2023/05/Bhorat-et.-al-May-2023-Digitalization-and-digital-skills-in-Africa-2.pdf",
        "Brevo. (n.d.). Transactional email API documentation. https://developers.brevo.com/docs/send-a-transactional-email",
        "E2B. (n.d.). Sandbox documentation. https://e2b.dev/docs",
        "Express.js. (n.d.). Express routing guide. https://expressjs.com/en/guide/routing.html",
        "GitHub Docs. (n.d.). REST API endpoints for repositories. https://docs.github.com/en/rest/repos/repos",
        "Google AI for Developers. (n.d.). Gemini API documentation. https://ai.google.dev/gemini-api/docs",
        "Hakizayezu, J., & Maniraho, J. (2022). Challenges facing technical and vocational education and training institutions on youth employment in Gasabo District, Rwanda. Journal of Research Innovation and Implications in Education, 6(3), 572-580. https://jriiejournal.com/wp-content/uploads/2022/10/JRIIE-6-3-050.pdf",
        "Newfarmer, R., & Twum, A. (2022). Employment creation potential, labour skills requirements, and skill gaps for young people: A Rwanda case study. Brookings Institution. https://www.brookings.edu/wp-content/uploads/2022/02/Rwanda-IWOSS.pdf",
        "OWASP Foundation. (n.d.). Authentication cheat sheet. https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html",
        "React. (n.d.). Learn React. https://react.dev/learn",
        "Rwanda Development Board. (2022). ICT sector labour market brief. https://rdb.rw/wp-content/uploads/2022/07/ICT-Sector-Labour-Market-Brief.pdf",
        "Rwanda Information Society Authority. (2023). National digital skills framework. https://www.risa.gov.rw/fileadmin/user_upload/RISA/Publications/4.Policies/National_Digital_Skills_Framework__NDSF_.pdf",
        "Rwanda Polytechnic. (2023). TVET competency-based curriculum development framework. https://www.rp.ac.rw/fileadmin/user_upload/RP/Publications/1._TVET_CBC_Development_Framework.pdf",
        "UNESCO-UNEVOC. (2017). Toolkits for TVET providers. https://connect.unevoc.unesco.org/home/Toolkits+for+TVET+providers",
        "World Bank Group. (2025, June 18). Transforming Rwanda's workforce: A skills-led approach for jobs and growth. https://www.worldbank.org/en/news/feature/2025/06/18/transforming-afe-rwanda-workforce-a-skills-led-approach-for-jobs-and-growth",
        "Zhong, Z., & Juwaheer, S. (2024). Digital competence development in TVET with a competency-based whole-institution approach. Vocation, Technology & Education, 1(2). https://doi.org/10.54844/vte.2024.0591",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(ref)
        r.font.name = "Arial"
        r.font.size = Pt(10)


def appendix(doc):
    heading(doc, "Appendix A: Missing Evidence Needed Before Full Submission", 1)
    for item in ["Signed declaration and supervisor certification.", "Final supervisor name.", "Actual user survey results from the proposed 50-120 ICT TVET graduates.", "Pre-assessment and post-assessment competency-awareness scores.", "Expert agreement percentage and usability rating.", "Final demo video link if the placeholder changes."]:
        bullet(doc, item)
    heading(doc, "Appendix B: Live Links", 1)
    paragraph(doc, "Frontend: https://rtb-graduate-project.vercel.app/")
    paragraph(doc, "Backend health endpoint: https://rtb-graduate-project.onrender.com/api/health")
    paragraph(doc, "Demo/video evidence placeholder: https://docs.google.com/document/d/1c0gMfJDdh6FZgHur0QYa_gu9NMmWGgFYGAZr0C8HBaY/edit?tab=t.0")


def main():
    architecture = create_architecture()
    workflow = create_diagram("workflow", "Core Assessment Workflow", [
        "User selected RTB-aligned ICT competency",
        "User submitted GitHub repository and theory answers",
        "System reviewed repository and scored theory evidence",
        "System calculated final score and compared benchmark",
        "System classified gap and generated Gemini recommendation",
        "System notified user and generated report",
    ])
    usecase = create_diagram("usecase", "Role-Based Use Cases", [
        "Learner registered, verified email, completed profile, and took assessment",
        "Learner viewed gap results, recommendations, reports, and notifications",
        "Organization admin managed organization users and organization results",
        "Admin managed users, organizations, competencies, benchmarks, and checklists",
        "System generated scores, gaps, recommendations, notifications, and reports",
    ])

    doc = Document(TEMPLATE)
    clear_document(doc)
    configure_document(doc)
    preliminaries(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc, architecture, workflow, usecase)
    chapter_four(doc)
    chapter_five(doc)
    chapter_six(doc)
    references(doc)
    appendix(doc)

    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()


